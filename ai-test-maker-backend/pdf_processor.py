import fitz  #type: ignore
from PIL import Image
import io
import numpy as np
import docx

class PDFProcessor:
    def __init__(self):
        self.page_samples = {
            10: (3, 5),
            25: (5, 10),
            50: (10, 20),
            100: (20, 40)
        }

    def process_file(self, file_path, marks):
        """Extract text and images from a file (PDF or DOCX)"""
        if file_path.lower().endswith('.pdf'):
            return self._process_pdf(file_path, marks)
        elif file_path.lower().endswith('.docx'):
            return self._process_docx(file_path)
        else:
            raise ValueError("Unsupported file type. Please provide a .pdf or .docx file.")

    def _process_pdf(self, pdf_path, marks):
        """Extract text and images from PDF based on marks"""
        doc = fitz.open(pdf_path)
        total_pages = len(doc)

        # Determine pages to sample
        min_pages, max_pages = self.page_samples[marks]
        pages_to_sample = min(max_pages, total_pages)
        pages_to_sample = max(min_pages, pages_to_sample)

        # Sample pages evenly
        if pages_to_sample >= total_pages:
            selected_pages = list(range(total_pages))
        else:
            step = total_pages / pages_to_sample
            selected_pages = [int(i * step) for i in range(pages_to_sample)]

        # Extract content
        text_content = []
        images = []

        for page_num in selected_pages:
            page = doc[page_num]

            # Extract text
            text = page.get_text()
            if text.strip():
                text_content.append(text)

            # Extract images
            image_list = page.get_images()
            for img_index, img in enumerate(image_list[:2]):  # Max 2 images per page
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]

                    # Convert to PIL Image and downscale
                    img_pil = Image.open(io.BytesIO(image_bytes))
                    img_pil = img_pil.convert('RGB')
                    img_pil.thumbnail((224, 224), Image.Resampling.LANCZOS)

                    # FIX: Remove the non-serializable numpy array. It is not used by the
                    # frontend and was causing the JSON serialization error.
                    images.append({
                        'page': page_num,
                        'index': img_index,
                        # 'array': np.array(img_pil), # This was the cause of the error
                        'size': img_pil.size
                    })
                except:
                    continue

        doc.close()

        # Chunk text into ~1.5k token segments (approx 6k chars)
        combined_text = "\n\n".join(text_content)
        chunks = self.chunk_text(combined_text, chunk_size=6000)

        return {
            'text_chunks': chunks,
            'images': images,
            'total_pages': total_pages,
            'sampled_pages': len(selected_pages)
        }

    def _process_docx(self, docx_path):
        """Extract text from a DOCX file"""
        try:
            document = docx.Document(docx_path)
            text_content = [p.text for p in document.paragraphs if p.text.strip()]
            combined_text = "\n\n".join(text_content)

            # Chunk the text using the existing method
            chunks = self.chunk_text(combined_text, chunk_size=6000)

            # DOCX processing does not currently support images or page sampling
            return {
                'text_chunks': chunks,
                'images': [],
                'total_pages': len(document.paragraphs) // 50, # Rough estimate for context
                'sampled_pages': 'N/A'
            }
        except Exception as e:
            print(f"Error processing DOCX file: {e}")
            return {'text_chunks': [], 'images': [], 'total_pages': 0, 'sampled_pages': 0}


    def chunk_text(self, text, chunk_size=6000):
        """Split text into chunks"""
        chunks = []
        current_chunk = ""

        # Split by paragraphs
        paragraphs = text.split('\n\n')

        for para in paragraphs:
            if len(current_chunk) + len(para) < chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"

        if current_chunk:
            chunks.append(current_chunk.strip())

        return chunks if chunks else [text[:chunk_size]]
